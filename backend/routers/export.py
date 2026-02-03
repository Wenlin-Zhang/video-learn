"""导出API路由"""
import logging
from io import BytesIO
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse, StreamingResponse

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from models.schemas import Lecture
from services.lecture_generator import get_lecture_generator

logger = logging.getLogger(__name__)
router = APIRouter()


def lecture_to_markdown(lecture: Lecture) -> str:
    """
    将讲义转换为Markdown格式
    
    Args:
        lecture: 讲义对象
        
    Returns:
        Markdown格式字符串
    """
    lines = []
    
    # 标题
    lines.append(f"# {lecture.title}")
    lines.append("")
    
    # 元信息
    lines.append(f"> 视频文件: {lecture.metadata.video_file}")
    lines.append(f"> 时长: {lecture.metadata.duration / 60:.1f} 分钟")
    lines.append(f"> 生成时间: {lecture.metadata.created_at}")
    lines.append("")
    lines.append("---")
    lines.append("")
    
    # 目录
    lines.append("## 目录")
    lines.append("")
    for section in lecture.sections:
        time_str = format_time(section.start_time)
        lines.append(f"- [{section.title}](#{section.id}) ({time_str})")
    lines.append("")
    lines.append("---")
    lines.append("")
    
    # 各小节内容
    for section in lecture.sections:
        start_str = format_time(section.start_time)
        end_str = format_time(section.end_time)
        
        lines.append(f"<a id='{section.id}'></a>")
        lines.append("")
        lines.append(f"## {section.id}. {section.title}")
        lines.append("")
        lines.append(f"*时间: {start_str} - {end_str}*")
        lines.append("")
        
        # 小节内容（已经是Markdown格式）
        # 去掉可能重复的标题
        content = section.content
        if content.startswith(f"## {section.title}"):
            content = content[len(f"## {section.title}"):].strip()
        
        lines.append(content)
        lines.append("")
        lines.append("---")
        lines.append("")
    
    return "\n".join(lines)


def format_time(seconds: float) -> str:
    """将秒数格式化为 MM:SS 格式"""
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{minutes:02d}:{secs:02d}"


def lecture_to_docx(lecture: Lecture) -> bytes:
    """
    将讲义转换为Word文档
    
    Args:
        lecture: 讲义对象
        
    Returns:
        Word文档的字节内容
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx库未安装，请运行: pip install python-docx"
        )
    
    doc = Document()
    
    # 标题
    title = doc.add_heading(lecture.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 元信息
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"视频文件: {lecture.metadata.video_file}\n")
    meta_para.add_run(f"时长: {lecture.metadata.duration / 60:.1f} 分钟\n")
    meta_para.add_run(f"生成时间: {lecture.metadata.created_at}")
    
    doc.add_paragraph()  # 空行
    
    # 目录
    doc.add_heading("目录", level=1)
    for section in lecture.sections:
        time_str = format_time(section.start_time)
        doc.add_paragraph(f"{section.id}. {section.title} ({time_str})", style="List Number")
    
    doc.add_page_break()
    
    # 各小节内容
    for section in lecture.sections:
        start_str = format_time(section.start_time)
        end_str = format_time(section.end_time)
        
        doc.add_heading(f"{section.id}. {section.title}", level=1)
        
        time_para = doc.add_paragraph()
        time_run = time_para.add_run(f"时间: {start_str} - {end_str}")
        time_run.italic = True
        
        doc.add_paragraph()
        
        # 处理Markdown内容，简单转换
        content = section.content
        # 去掉可能重复的标题
        if content.startswith(f"## {section.title}"):
            content = content[len(f"## {section.title}"):].strip()
        
        # 按段落分割
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                # 处理列表项
                if para_text.strip().startswith("- "):
                    for line in para_text.split("\n"):
                        if line.strip().startswith("- "):
                            doc.add_paragraph(line.strip()[2:], style="List Bullet")
                        elif line.strip():
                            doc.add_paragraph(line.strip())
                elif para_text.strip().startswith("1. ") or para_text.strip().startswith("1. "):
                    for line in para_text.split("\n"):
                        if line.strip() and line.strip()[0].isdigit():
                            # 去掉数字和点
                            text = line.strip()
                            if ". " in text:
                                text = text.split(". ", 1)[1]
                            doc.add_paragraph(text, style="List Number")
                        elif line.strip():
                            doc.add_paragraph(line.strip())
                else:
                    # 普通段落，去掉Markdown格式
                    clean_text = para_text.replace("**", "").replace("*", "").replace("`", "")
                    doc.add_paragraph(clean_text.strip())
        
        doc.add_paragraph()  # 空行
    
    # 保存到字节流
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@router.get("/markdown/{task_id}")
async def export_markdown(task_id: str):
    """
    导出Markdown格式讲义
    
    Args:
        task_id: 任务ID
        
    Returns:
        Markdown文件
    """
    from routers.video import processing_tasks
    
    if task_id not in processing_tasks:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    task = processing_tasks[task_id]
    if task["status"] != "completed":
        raise HTTPException(status_code=400, detail="任务尚未完成")
    
    # 加载讲义
    lecture_gen = get_lecture_generator()
    lecture = lecture_gen.load_lecture(task["result"]["lecture_path"])
    
    # 转换为Markdown
    md_content = lecture_to_markdown(lecture)
    
    # 保存文件
    output_path = Path(task["result"]["lecture_path"]).with_suffix(".md")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    
    return FileResponse(
        path=str(output_path),
        filename=f"{lecture.title}.md",
        media_type="text/markdown"
    )


@router.get("/word/{task_id}")
async def export_word(task_id: str):
    """
    导出Word格式讲义
    
    Args:
        task_id: 任务ID
        
    Returns:
        Word文档
    """
    from routers.video import processing_tasks
    
    if task_id not in processing_tasks:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    task = processing_tasks[task_id]
    if task["status"] != "completed":
        raise HTTPException(status_code=400, detail="任务尚未完成")
    
    # 加载讲义
    lecture_gen = get_lecture_generator()
    lecture = lecture_gen.load_lecture(task["result"]["lecture_path"])
    
    # 转换为Word
    docx_content = lecture_to_docx(lecture)
    
    # 保存文件
    output_path = Path(task["result"]["lecture_path"]).with_suffix(".docx")
    with open(output_path, "wb") as f:
        f.write(docx_content)
    
    return FileResponse(
        path=str(output_path),
        filename=f"{lecture.title}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.post("/from-lecture")
async def export_from_lecture_file(
    lecture_path: str,
    format: str = "markdown"
):
    """
    从讲义文件导出
    
    Args:
        lecture_path: 讲义JSON文件路径
        format: 导出格式 (markdown 或 word)
        
    Returns:
        导出的文件
    """
    lecture_path = Path(lecture_path)
    
    if not lecture_path.exists():
        raise HTTPException(status_code=404, detail="讲义文件不存在")
    
    # 加载讲义
    lecture_gen = get_lecture_generator()
    lecture = lecture_gen.load_lecture(str(lecture_path))
    
    if format == "markdown":
        md_content = lecture_to_markdown(lecture)
        output_path = lecture_path.with_suffix(".md")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        
        return FileResponse(
            path=str(output_path),
            filename=f"{lecture.title}.md",
            media_type="text/markdown"
        )
    
    elif format == "word":
        docx_content = lecture_to_docx(lecture)
        output_path = lecture_path.with_suffix(".docx")
        with open(output_path, "wb") as f:
            f.write(docx_content)
        
        return FileResponse(
            path=str(output_path),
            filename=f"{lecture.title}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    else:
        raise HTTPException(status_code=400, detail=f"不支持的格式: {format}")
